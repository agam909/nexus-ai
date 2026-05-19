"""RAG pipeline backed by Qdrant Cloud + HuggingFace Inference API embeddings.

Per-user data isolation is enforced by storing `user_id` in every vector's
payload and filtering on it during retrieval.
"""
from __future__ import annotations

import uuid
from functools import lru_cache
from pathlib import Path
from typing import Iterable, List, Optional

from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from qdrant_client.http import models as qm

from .config import get_settings
from .embeddings import get_embedder

# ───────────────────────── document loaders ─────────────────────────


def _load_text(path: str) -> List[LCDocument]:
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        from pypdf import PdfReader  # local import; tiny dep
        reader = PdfReader(str(p))
        out: List[LCDocument] = []
        for i, page in enumerate(reader.pages):
            txt = (page.extract_text() or "").strip()
            if txt:
                out.append(LCDocument(page_content=txt, metadata={"page": i}))
        return out
    if ext == ".docx":
        import docx  # python-docx
        d = docx.Document(str(p))
        text = "\n".join(par.text for par in d.paragraphs if par.text.strip())
        return [LCDocument(page_content=text, metadata={})] if text else []
    # txt / md / anything text-ish
    try:
        text = p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = p.read_text(encoding="latin-1", errors="ignore")
    return [LCDocument(page_content=text, metadata={})] if text.strip() else []


def load_and_split(path: str, doc_id: str, file_name: str) -> List[LCDocument]:
    s = get_settings()
    raw = _load_text(path)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=s.chunk_size,
        chunk_overlap=s.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(raw)
    for i, c in enumerate(chunks):
        c.metadata = {
            **(c.metadata or {}),
            "doc_id": doc_id,
            "file_name": file_name,
            "chunk_index": i,
        }
    return chunks


# ───────────────────────── Qdrant client ────────────────────────────


@lru_cache
def get_qdrant() -> QdrantClient:
    s = get_settings()
    if not s.qdrant_url:
        raise RuntimeError(
            "QDRANT_URL is not set. Create a free cluster at "
            "https://cloud.qdrant.io and add QDRANT_URL + QDRANT_API_KEY to .env"
        )
    client = QdrantClient(
        url=s.qdrant_url,
        api_key=s.qdrant_api_key or None,
        prefer_grpc=False,
        timeout=30.0,
    )
    return client


def ensure_collection() -> None:
    """Idempotently create the collection if it doesn't exist."""
    s = get_settings()
    client = get_qdrant()
    existing = {c.name for c in client.get_collections().collections}
    if s.qdrant_collection in existing:
        return
    client.create_collection(
        collection_name=s.qdrant_collection,
        vectors_config=qm.VectorParams(
            size=s.embedding_dim,
            distance=qm.Distance.COSINE,
        ),
    )


# ───────────────────────── ingest / delete / retrieve ───────────────


async def ingest(path: str, doc_id: str, file_name: str, user_id: str) -> int:
    chunks = load_and_split(path, doc_id, file_name)
    if not chunks:
        return 0

    s = get_settings()
    ensure_collection()

    embedder = get_embedder()
    texts = [c.page_content for c in chunks]
    vectors = await embedder.aembed_documents(texts)

    points: List[qm.PointStruct] = []
    for c, v in zip(chunks, vectors):
        meta = c.metadata or {}
        points.append(
            qm.PointStruct(
                id=uuid.uuid4().hex,
                vector=v,
                payload={
                    "user_id": user_id,
                    "doc_id": doc_id,
                    "file_name": file_name,
                    "chunk_index": meta.get("chunk_index", 0),
                    "page": meta.get("page"),
                    "text": c.page_content,
                },
            )
        )

    client = get_qdrant()
    client.upsert(collection_name=s.qdrant_collection, points=points, wait=True)
    return len(chunks)


def delete_doc(doc_id: str, user_id: str) -> None:
    s = get_settings()
    try:
        ensure_collection()
        client = get_qdrant()
        client.delete(
            collection_name=s.qdrant_collection,
            points_selector=qm.FilterSelector(
                filter=qm.Filter(
                    must=[
                        qm.FieldCondition(key="doc_id", match=qm.MatchValue(value=doc_id)),
                        qm.FieldCondition(key="user_id", match=qm.MatchValue(value=user_id)),
                    ]
                )
            ),
            wait=True,
        )
    except Exception:
        pass


async def retrieve(query: str, user_id: str, k: Optional[int] = None) -> List[LCDocument]:
    s = get_settings()
    ensure_collection()
    embedder = get_embedder()
    vec = await embedder.aembed_query(query)
    client = get_qdrant()
    hits = client.search(
        collection_name=s.qdrant_collection,
        query_vector=vec,
        limit=k or s.top_k,
        query_filter=qm.Filter(
            must=[
                qm.FieldCondition(key="user_id", match=qm.MatchValue(value=user_id))
            ]
        ),
    )
    out: List[LCDocument] = []
    for h in hits:
        p = h.payload or {}
        out.append(
            LCDocument(
                page_content=p.get("text", ""),
                metadata={
                    "doc_id": p.get("doc_id"),
                    "file_name": p.get("file_name", "unknown"),
                    "page": p.get("page"),
                    "score": h.score,
                },
            )
        )
    return out


def collection_count(user_id: str) -> int:
    s = get_settings()
    try:
        ensure_collection()
        client = get_qdrant()
        res = client.count(
            collection_name=s.qdrant_collection,
            count_filter=qm.Filter(
                must=[
                    qm.FieldCondition(key="user_id", match=qm.MatchValue(value=user_id))
                ]
            ),
            exact=False,
        )
        return res.count
    except Exception:
        return 0
